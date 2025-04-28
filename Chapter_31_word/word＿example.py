# word_example.py

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement

path = '/Users/lei/Downloads/示例文檔.docx'
name = 'XX市ABC有限责任公司'
date = '202X年12月31日'
period = '202X年度'


# 創建一個文檔實例
doc = Document()


# 添加页眉
header = doc.sections[0].header                                 # sections[0]是文檔的第一個節
header_paragraph = header.paragraphs[0]                         # 頁眉的第一個段落
header_paragraph.text = '这是页眉'                               # 設置頁眉的內容 
header_run = header_paragraph.runs[0]                           # 第一个文本运行，每个Run是一段连续的、具有相同格式的文本
header_run.font.name = 'SimSun'                                 # 设置字体：仅影响拉丁字符字体
header_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')  # 设置中文字体
header_run.font.size = Pt(12)                                   # 字號
header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER      # 居中

# 設置頁眉帶橫線
header_p = header_paragraph._element                            # 获取段落的底层 XML 元素
header_pPr = header_p.get_or_add_pPr()                          # 获取段落属性
header_pbdr = OxmlElement('w:pBdr')                             # 创建边框元素
header_bottom = OxmlElement('w:bottom')                         # 添加底部边框
header_bottom.set(qn('w:val'), 'single')                        # 线型（single 表示单线）
header_bottom.set(qn('w:sz'), '8')                              # 线条宽度（1/8 磅）
header_bottom.set(qn('w:space'), '4')                           # 与文本的间距
header_bottom.set(qn('w:color'), '000000')                      # 颜色（黑色）
header_pbdr.append(header_bottom)                               # 将底部边框添加到段落边框中
header_pPr.append(header_pbdr)                                  # 将边框设置应用到段落

# 首页与后续页不同： 如果文档要求首页与其他页面的页眉不同，可以启用：
# doc.sections[0].different_first_page_header_footer = True

# 调试 XML：可以使用.xml属性检查底层XML结构，例如：
# print(paragraph._element.xml)


# 添加一個標題，級別為0-9
# 每個標題級別會對應Word的預設樣式（例如字體大小和格式），但可以透過python-docx的樣式功能調整，以自訂樣式。
doc.add_heading('這是標題 0', 0)
doc.add_heading('這是標題 0', level=0)
doc.add_heading('這是標題 1', level=1)
doc.add_heading('這是標題 2', level=2)
doc.add_heading('這是標題 3', level=3)
doc.add_heading('這是標題 4', level=4)
doc.add_heading('這是標題 5', level=5)
doc.add_heading('這是標題 6', level=6)
doc.add_heading('這是標題 7', level=7)
doc.add_heading('這是標題 8', level=8)
doc.add_heading('這是標題 9', level=9)


# 添加段落
paragraph = doc.add_paragraph('這是一個段落的示例文本。')
run = paragraph.runs[0]
run.font.size = Pt(12)                                  # 給段落的文字設置字體大小

# 新增一個居中對齊的段落
paragraph = doc.add_paragraph('這個段落是居中對齊的。')
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER     # 居中
# 對齊方式
# WD_PARAGRAPH_ALIGNMENT.LEFT       # 左对齐
# WD_PARAGRAPH_ALIGNMENT.CENTER     # 居中对齐
# WD_PARAGRAPH_ALIGNMENT.RIGHT      # 右对齐
# WD_PARAGRAPH_ALIGNMENT.JUSTIFY    # 两端对齐
# WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE # 分散对齐


# 插入分頁符
doc.add_page_break()


# 再添加一個段落
doc.add_paragraph('新頁面的段落。')

# 添加段落并设置字体和对齐方式
paragraph_example = doc.add_paragraph()
text_example = f'    我们{name}（以下简称“贵公司”）财务报表，包括{date}的资产负债表，{period}的利润表、现金流量表、所有者权益变动表以及相关财务报表附注。'
run_example = paragraph_example.add_run(text_example)
run_example.font.name = 'SimSun'                                    # 设置字体
run_example._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')     # 设置中文字体
run_example.font.size = Pt(12)                                      # 设置字号
run_example.bold = True                                             # 加粗
paragraph_example.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT           # 设置左对齐


# 添加表格一個4x3的表格
table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'                                                  # 设置表格样式（默認樣式）

# 填充表格内容并设置对齐方式
for i, row in enumerate(table.rows):
    for j, cell in enumerate(row.cells):
        cell.text = ''                                                      # 先清空默认文本
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(f'单元格({i+1}, {j+1})')
 
        run.font.name = 'SimSun'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        run.font.size = Pt(12)
        if i == 0:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER    # 第一行居中对齐
        else:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT      # 其他行靠右对齐


# 添加页脚
footer = doc.sections[0].footer
footer_paragraph = footer.paragraphs[0]
footer_paragraph.text = '这是页脚'
footer_run = footer_paragraph.runs[0]
footer_run.font.name = 'SimSun'
footer_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
footer_run.font.size = Pt(12)
footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# 設置頁眉帶橫線
footer_p = footer_paragraph._element                            # 获取段落的底层 XML 元素
footer_pPr = footer_p.get_or_add_pPr()                          # 获取段落属性
footer_pbdr = OxmlElement('w:pBdr')                             # 创建边框元素
footer_top = OxmlElement('w:top')                               # 添加底部边框
footer_top.set(qn('w:val'), 'single')                           # 线型（single 表示单线）
footer_top.set(qn('w:sz'), '8')                                 # 线条宽度（1/8 磅）
footer_top.set(qn('w:space'), '4')                              # 与文本的间距
footer_top.set(qn('w:color'), '000000')                         # 颜色（黑色）
footer_pbdr.append(footer_top)                                  # 将底部边框添加到段落边框中
footer_pPr.append(footer_pbdr)                                  # 将边框设置应用到段落


# 保存文檔
doc.save(path)